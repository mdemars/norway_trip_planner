"""Trip export helpers — generate Word, PDF, and Excel documents from trip data."""

import calendar
import json
import os
import urllib.request
import urllib.parse
from datetime import date
from io import BytesIO

# ── Colour palette (8 cycling stop colours, light variants for backgrounds) ──
STOP_COLORS = [
    '#4285F4', '#DB4437', '#0F9D58', '#F4B400',
    '#AB47BC', '#00ACC1', '#FF7043', '#9CCC65',
]
STOP_COLORS_LIGHT = [
    '#C6D9FB', '#F7CBCA', '#C0E8D0', '#FDE9A2',
    '#E8D5F0', '#B2EAF0', '#FFD0C0', '#DFF0C4',
]

_MONTH_KEYS = ['january', 'february', 'march', 'april', 'may', 'june',
               'july', 'august', 'september', 'october', 'november', 'december']
# calendar.monthcalendar weeks start on Monday → order is Mon..Sun
_WEEKDAY_KEYS = ['mon', 'tue', 'wed', 'thu', 'fri', 'sat', 'sun']


# ── Translation helper ────────────────────────────────────────────────────────

class Translations:
    """Thin wrapper around a locale JSON file with dot-notation access."""

    _SUPPORTED = {'en', 'fr', 'de'}

    def __init__(self, lang: str = 'en'):
        self.lang = lang if lang in self._SUPPORTED else 'en'
        path = os.path.join(os.path.dirname(__file__),
                            'static', 'locales', self.lang, 'translation.json')
        with open(path, 'r', encoding='utf-8') as fh:
            self._data = json.load(fh)

    def t(self, key: str, default: str = '') -> str:
        """Translate a dot-separated key, returning *default* when missing."""
        node = self._data
        for part in key.split('.'):
            if isinstance(node, dict):
                node = node.get(part, {})
            else:
                return default
        return node if isinstance(node, str) else default

    def month_name(self, month_num: int) -> str:
        return self.t(f'calendar.months.{_MONTH_KEYS[month_num - 1]}',
                      calendar.month_name[month_num])

    def weekday_abbr(self, idx: int) -> str:
        """idx: 0=Mon … 6=Sun (matching calendar.monthcalendar order)."""
        return self.t(f'calendar.weekdays.{_WEEKDAY_KEYS[idx]}',
                      _WEEKDAY_KEYS[idx].capitalize())

    def fmt_date(self, dt) -> str:
        if dt is None:
            return ''
        month = self.month_name(dt.month)
        if self.lang == 'de':
            return f'{dt.day}. {month} {dt.year}'
        return f'{dt.day} {month} {dt.year}'

    def nights_label(self, n: int) -> str:
        word = self.t('export.night') if n == 1 else self.t('export.nights')
        return f'{n} {word}'


# ── Shared helpers ────────────────────────────────────────────────────────────

def _date_range_label(stops, tr: Translations) -> str:
    dated = [s for s in stops if s.start_date]
    if not dated:
        return ''
    first = min(s.start_date for s in dated)
    last = max((s.end_date or s.start_date) for s in dated)
    return f'{tr.fmt_date(first)} – {tr.fmt_date(last)}'


def _nights(stop):
    if stop.start_date and stop.end_date:
        return (stop.end_date - stop.start_date).days
    return None


def _calendar_months(stops):
    dated = [s for s in stops if s.start_date]
    if not dated:
        return []
    start = min(s.start_date for s in dated).date()
    end = max((s.end_date or s.start_date) for s in dated).date()
    months = []
    cur = date(start.year, start.month, 1)
    while cur <= end:
        months.append((cur.year, cur.month))
        cur = date(cur.year + (cur.month == 12), (cur.month % 12) + 1, 1)
    return months


def _stop_index_for_date(d, stops):
    for i, stop in enumerate(stops):
        if stop.start_date and stop.end_date:
            if stop.start_date.date() <= d <= stop.end_date.date():
                return i
    return None


def _get_route_polylines(stops, trip):
    """Return a list of encoded polyline strings by calling the Directions API.
    Falls back to an empty list on any error so the caller can degrade gracefully.
    """
    from services import route_service

    route_points = []
    if trip.start_location_latitude and trip.start_location_longitude:
        route_points.append({
            'id': 0,
            'name': trip.start_location_address or 'Start',
            'latitude': trip.start_location_latitude,
            'longitude': trip.start_location_longitude,
        })
    for s in stops:
        if s.latitude and s.longitude:
            route_points.append({
                'id': s.id,
                'name': s.name,
                'latitude': s.latitude,
                'longitude': s.longitude,
            })
    if trip.end_location_latitude and trip.end_location_longitude:
        route_points.append({
            'id': -1,
            'name': trip.end_location_address or 'End',
            'latitude': trip.end_location_latitude,
            'longitude': trip.end_location_longitude,
        })

    if len(route_points) < 2:
        return []

    try:
        route_data = route_service.calculate_route(route_points)
        return [seg['polyline'] for seg in route_data.get('segments', [])
                if seg.get('polyline')]
    except Exception:
        return []


def _get_static_map_image(stops, trip, api_key, width=800, height=500):
    if not api_key:
        return None

    parts = [f'size={width}x{height}', 'maptype=roadmap']

    # Prefer road-following polylines from the Directions API
    polylines = _get_route_polylines(stops, trip)
    if polylines:
        for pl in polylines:
            parts.append(
                f'path=color:0x4285F4|weight:3|enc:{urllib.parse.quote(pl)}')
    else:
        # Fallback: straight lines through stop coordinates
        coords = []
        if trip.start_location_latitude and trip.start_location_longitude:
            coords.append((trip.start_location_latitude, trip.start_location_longitude))
        for s in stops:
            if s.latitude and s.longitude:
                coords.append((s.latitude, s.longitude))
        if trip.end_location_latitude and trip.end_location_longitude:
            coords.append((trip.end_location_latitude, trip.end_location_longitude))
        if not coords:
            return None
        if len(coords) > 50:
            step = max(1, len(coords) // 50)
            coords = coords[::step]
        path_points = '|'.join(f'{lat:.4f},{lng:.4f}' for lat, lng in coords)
        parts.append(f'path=color:0x4285F4|weight:3|{urllib.parse.quote(path_points)}')

    # Markers for each stop
    for i, s in enumerate(stops):
        if s.latitude and s.longitude:
            label = str(i + 1) if i < 9 else 'S'
            color = STOP_COLORS[i % len(STOP_COLORS)].lstrip('#')
            parts.append(
                f'markers={urllib.parse.quote(f"color:0x{color}|label:{label}|{s.latitude:.4f},{s.longitude:.4f}")}')

    parts.append(f'key={api_key}')
    url = 'https://maps.googleapis.com/maps/api/staticmap?' + '&'.join(parts)
    try:
        with urllib.request.urlopen(url, timeout=30) as resp:
            return resp.read()
    except Exception:
        return None


# ── Excel ─────────────────────────────────────────────────────────────────────

def generate_excel(trip, stops, bookmarks, lang='en'):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    tr = Translations(lang)

    wb = openpyxl.Workbook()

    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill('solid', fgColor='2E4057')
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    wrap = Alignment(wrap_text=True, vertical='top')
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _header_row(ws, headers):
        ws.row_dimensions[1].height = 22
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=col, value=h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = center
            c.border = border

    def _auto_width(ws, min_w=10, max_w=60):
        for col in ws.columns:
            length = max((len(str(cell.value or '')) for cell in col), default=0)
            ws.column_dimensions[get_column_letter(col[0].column)].width = \
                max(min_w, min(length + 4, max_w))

    def _style_data_rows(ws):
        alt = PatternFill('solid', fgColor='F5F7FA')
        for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            for cell in row:
                cell.alignment = wrap
                cell.border = border
                if row_idx % 2 == 0:
                    cell.fill = alt

    c = tr.t  # shorthand

    # ---- Sheet 1: Stops ----
    ws1 = wb.active
    ws1.title = c('export.stops', 'Stops')
    ws1.freeze_panes = 'A2'
    _header_row(ws1, [
        c('export.cols.num', '#'),
        c('export.cols.name', 'Name'),
        c('export.cols.address', 'Address'),
        c('export.cols.startDate', 'Start Date'),
        c('export.cols.endDate', 'End Date'),
        c('export.cols.nights', 'Nights'),
        c('export.cols.distanceKm', 'Distance (km)'),
        c('export.cols.description', 'Description'),
        c('export.cols.url', 'URL'),
    ])
    for i, s in enumerate(stops, 1):
        ws1.append([
            i, s.name, s.address or '',
            tr.fmt_date(s.start_date),
            tr.fmt_date(s.end_date),
            _nights(s),
            round(s.distance_km, 1) if s.distance_km else '',
            s.description or '',
            s.url or '',
        ])
    _style_data_rows(ws1)
    _auto_width(ws1)

    # ---- Sheet 2: Activities ----
    ws2 = wb.create_sheet(c('export.activities', 'Activities'))
    ws2.freeze_panes = 'A2'
    _header_row(ws2, [
        c('export.cols.stopNum', 'Stop #'),
        c('export.cols.stopName', 'Stop Name'),
        c('export.cols.activity', 'Activity'),
        c('export.cols.description', 'Description'),
        c('export.cols.url', 'URL'),
    ])
    for i, s in enumerate(stops, 1):
        for act in s.activities:
            ws2.append([i, s.name, act.name, act.description or '', act.url or ''])
    _style_data_rows(ws2)
    _auto_width(ws2)

    # ---- Sheet 3: Bookmarks ----
    ws3 = wb.create_sheet(c('bookmarks.title', 'Bookmarks'))
    ws3.freeze_panes = 'A2'
    _header_row(ws3, [
        c('export.cols.description', 'Description'),
        c('export.cols.url', 'URL'),
    ])
    for bm in bookmarks:
        ws3.append([bm.description or '', bm.url or ''])
    _style_data_rows(ws3)
    _auto_width(ws3)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ── Word ──────────────────────────────────────────────────────────────────────

def generate_word(trip, stops, bookmarks, api_key=None, lang='en'):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tr = Translations(lang)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color.lstrip('#'))
        tcPr.append(shd)

    def _remove_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _para(text, size=None, bold=False, color=None, align=None,
              space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if align:
            p.alignment = align
        run = p.add_run(text)
        if size:
            run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # ═══════════════════════════════════════
    # PAGE 1 — OVERVIEW
    # ═══════════════════════════════════════
    _para(trip.name, size=22, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    subtitle_parts = []
    dr = _date_range_label(stops, tr)
    if dr:
        subtitle_parts.append(dr)
    if trip.total_distance_km:
        subtitle_parts.append(f'{trip.total_distance_km:,.0f} km {tr.t("export.total", "total")}')
    if subtitle_parts:
        _para(' · '.join(subtitle_parts), size=11, color=(100, 100, 100),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 2-column overview table
    overview_tbl = doc.add_table(rows=1, cols=2)
    overview_tbl.style = 'Table Grid'
    overview_tbl.columns[0].width = Cm(9)
    overview_tbl.columns[1].width = Cm(10)
    left_cell = overview_tbl.cell(0, 0)
    right_cell = overview_tbl.cell(0, 1)
    for cell in [left_cell, right_cell]:
        _remove_cell_borders(cell)

    # Left: stop list
    left_cell.paragraphs[0].clear()
    run = left_cell.paragraphs[0].add_run(tr.t('export.stops', 'Stops'))
    run.bold = True
    run.font.size = Pt(12)

    for i, s in enumerate(stops):
        hex_c = STOP_COLORS[i % len(STOP_COLORS)]
        r, g, b = int(hex_c[1:3], 16), int(hex_c[3:5], 16), int(hex_c[5:7], 16)
        p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        dot = p.add_run(f'● {i + 1}. ')
        dot.font.color.rgb = RGBColor(r, g, b)
        dot.font.size = Pt(10)
        p.add_run(s.name).font.size = Pt(10)
        meta_parts = []
        if s.start_date:
            meta_parts.append(tr.fmt_date(s.start_date))
        n = _nights(s)
        if n is not None:
            meta_parts.append(tr.nights_label(n))
        if s.distance_km:
            meta_parts.append(f'{s.distance_km:.0f} km')
        if meta_parts:
            meta = p.add_run(f'\n  {" · ".join(meta_parts)}')
            meta.font.size = Pt(8)
            meta.font.color.rgb = RGBColor(120, 120, 120)

    # Right: calendar
    right_cell.paragraphs[0].clear()
    run = right_cell.paragraphs[0].add_run(tr.t('export.calendar', 'Calendar'))
    run.bold = True
    run.font.size = Pt(12)

    day_names = [tr.weekday_abbr(i) for i in range(7)]
    for year, month in _calendar_months(stops):
        mhd = right_cell.add_paragraph()
        mhd.paragraph_format.space_before = Pt(8)
        mhd.paragraph_format.space_after = Pt(2)
        mrun = mhd.add_run(f'{tr.month_name(month)} {year}')
        mrun.bold = True
        mrun.font.size = Pt(9)

        cal_tbl = right_cell.add_table(rows=1, cols=7)
        for col, dn in enumerate(day_names):
            hc = cal_tbl.cell(0, col)
            hc.text = dn
            hc.paragraphs[0].runs[0].font.size = Pt(7)
            hc.paragraphs[0].runs[0].bold = True
            hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(hc, 'E8EAED')
            hc.width = Cm(1.1)

        for week in calendar.monthcalendar(year, month):
            row = cal_tbl.add_row()
            for col, day_num in enumerate(week):
                cell = row.cells[col]
                cell.width = Cm(1.1)
                if day_num == 0:
                    _set_cell_bg(cell, 'F8F9FA')
                else:
                    d = date(year, month, day_num)
                    stop_idx = _stop_index_for_date(d, stops)
                    cell.text = str(day_num)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(7)
                    if stop_idx is not None:
                        _set_cell_bg(cell, STOP_COLORS_LIGHT[stop_idx % len(STOP_COLORS_LIGHT)])

    doc.add_page_break()

    # ═══════════════════════════════════════
    # DETAILED ITINERARY
    # ═══════════════════════════════════════
    doc.add_heading(tr.t('export.detailedItinerary', 'Detailed Itinerary'), level=1)

    for i, s in enumerate(stops):
        hex_c = STOP_COLORS[i % len(STOP_COLORS)]
        r, g, b = int(hex_c[1:3], 16), int(hex_c[3:5], 16), int(hex_c[5:7], 16)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(f'{i + 1}. ')
        num.font.size = Pt(13)
        num.bold = True
        num.font.color.rgb = RGBColor(r, g, b)
        title = p.add_run(s.name)
        title.font.size = Pt(13)
        title.bold = True

        meta_parts = []
        if s.start_date:
            if s.end_date:
                meta_parts.append(f'{tr.fmt_date(s.start_date)} – {tr.fmt_date(s.end_date)}')
            else:
                meta_parts.append(tr.fmt_date(s.start_date))
        n = _nights(s)
        if n is not None:
            meta_parts.append(tr.nights_label(n))
        if s.distance_km:
            meta_parts.append(f'{s.distance_km:.0f} {tr.t("export.kmFromPrevious", "km from previous")}')
        if meta_parts:
            _para(' · '.join(meta_parts), size=9, color=(100, 100, 100), space_after=4)

        if s.address:
            _para(s.address, size=9, color=(130, 130, 130), space_after=4)
        if s.description:
            _para(s.description, size=10, space_after=4)
        if s.url:
            _para(s.url, size=9, color=(66, 133, 244), space_after=4)

        if s.activities:
            _para(f'{tr.t("export.activities", "Activities")}:', size=10, bold=True, space_after=2)
            for act in s.activities:
                ap = doc.add_paragraph(style='List Bullet')
                ap.paragraph_format.space_after = Pt(2)
                ap.add_run(act.name).font.size = Pt(10)
                if act.description:
                    ap.add_run(f' — {act.description}').font.size = Pt(9)

    # ---- Bookmarks ----
    if bookmarks:
        doc.add_page_break()
        doc.add_heading(tr.t('export.bookmarksResources', 'Bookmarks & Resources'), level=1)
        for bm in bookmarks:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(3)
            if bm.description:
                bp.add_run(f'{bm.description}: ').font.size = Pt(10)
            url_run = bp.add_run(bm.url)
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(66, 133, 244)

    # ═══════════════════════════════════════
    # MAP PAGE
    # ═══════════════════════════════════════
    doc.add_page_break()
    doc.add_heading(tr.t('export.routeMap', 'Route Map'), level=1)

    map_bytes = _get_static_map_image(stops, trip, api_key)
    if map_bytes:
        doc.add_picture(BytesIO(map_bytes), width=Inches(6.5))
    else:
        _para(tr.t('export.mapUnavailable', 'Map unavailable — no coordinates found.'),
              size=10, color=(150, 150, 150))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF ───────────────────────────────────────────────────────────────────────

def generate_pdf(trip, stops, bookmarks, api_key=None, lang='en'):
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                     TableStyle, Image as RLImage, PageBreak, HRFlowable)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4

    tr = Translations(lang)

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=2*cm, bottomMargin=2*cm,
                             leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()
    W = A4[0] - 5*cm

    title_style = ParagraphStyle('TripTitle', parent=styles['Title'],
                                  fontSize=22, spaceAfter=6, alignment=TA_CENTER)
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'],
                                     fontSize=11,
                                     textColor=colors.HexColor('#646464'),
                                     spaceAfter=20, alignment=TA_CENTER)
    h1_style = ParagraphStyle('H1', parent=styles['Heading1'],
                               fontSize=16, textColor=colors.HexColor('#2E4057'),
                               spaceBefore=12, spaceAfter=6)
    stop_title_style = ParagraphStyle('StopTitle', parent=styles['Normal'],
                                       fontSize=13, fontName='Helvetica-Bold',
                                       spaceBefore=12, spaceAfter=2)
    meta_style = ParagraphStyle('Meta', parent=styles['Normal'],
                                  fontSize=9,
                                  textColor=colors.HexColor('#646464'),
                                  spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                  fontSize=10, spaceAfter=4)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'],
                                    fontSize=10, leftIndent=14,
                                    firstLineIndent=-14, spaceAfter=2)
    url_style = ParagraphStyle('URL', parent=styles['Normal'],
                                 fontSize=9,
                                 textColor=colors.HexColor('#4285F4'),
                                 spaceAfter=4)

    story = []

    # ═══════════════════════════════════════
    # PAGE 1 — OVERVIEW
    # ═══════════════════════════════════════
    story.append(Paragraph(trip.name, title_style))

    subtitle_parts = []
    dr = _date_range_label(stops, tr)
    if dr:
        subtitle_parts.append(dr)
    if trip.total_distance_km:
        subtitle_parts.append(
            f'{trip.total_distance_km:,.0f} km {tr.t("export.total", "total")}')
    if subtitle_parts:
        story.append(Paragraph(' · '.join(subtitle_parts), subtitle_style))

    # Stop list column
    stop_lines = [Paragraph(f'<b>{tr.t("export.stops", "Stops")}</b>',
                             ParagraphStyle('SLH', fontSize=11, spaceAfter=6))]
    for i, s in enumerate(stops):
        color = STOP_COLORS[i % len(STOP_COLORS)]
        label = f'<font color="{color}"><b>{i+1}.</b></font> <b>{s.name}</b>'
        meta_parts = []
        if s.start_date:
            meta_parts.append(tr.fmt_date(s.start_date))
        n = _nights(s)
        if n is not None:
            meta_parts.append(tr.nights_label(n))
        if s.distance_km:
            meta_parts.append(f'{s.distance_km:.0f} km')
        if meta_parts:
            label += f'<br/><font size="8" color="#888888">&nbsp;&nbsp;{" · ".join(meta_parts)}</font>'
        stop_lines.append(Paragraph(label,
                                     ParagraphStyle(f'SL{i}', fontSize=10,
                                                    spaceAfter=4, leading=14)))

    # Calendar column
    cal_content = [Paragraph(f'<b>{tr.t("export.calendar", "Calendar")}</b>',
                              ParagraphStyle('CLH', fontSize=11, spaceAfter=6))]
    day_names = [tr.weekday_abbr(i) for i in range(7)]

    for year, month in _calendar_months(stops):
        cal_content.append(Paragraph(
            f'<b>{tr.month_name(month)} {year}</b>',
            ParagraphStyle('MH', fontSize=9, spaceBefore=8, spaceAfter=2)))

        cal_weeks = calendar.monthcalendar(year, month)
        table_data = [day_names[:]]
        for week in cal_weeks:
            table_data.append([str(d) if d else '' for d in week])

        col_w = 1.1*cm
        cal_tbl = Table(table_data,
                         colWidths=[col_w]*7,
                         rowHeights=[0.55*cm] * (len(cal_weeks) + 1))
        style_cmds = [
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8EAED')),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#CCCCCC')),
        ]
        for r_idx, week in enumerate(cal_weeks, start=1):
            for c_idx, day_num in enumerate(week):
                if day_num:
                    d = date(year, month, day_num)
                    idx = _stop_index_for_date(d, stops)
                    if idx is not None:
                        style_cmds.append(
                            ('BACKGROUND', (c_idx, r_idx), (c_idx, r_idx),
                             colors.HexColor(STOP_COLORS_LIGHT[idx % len(STOP_COLORS_LIGHT)])))
        cal_tbl.setStyle(TableStyle(style_cmds))
        cal_content.append(cal_tbl)
        cal_content.append(Spacer(1, 2))

    overview_data = [[stop_lines, cal_content]]
    overview_tbl = Table(overview_data, colWidths=[W*0.45, W*0.55])
    overview_tbl.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#DDDDDD')),
        ('LINEAFTER', (0, 0), (0, -1), 0.5, colors.HexColor('#DDDDDD')),
    ]))
    story.append(overview_tbl)
    story.append(PageBreak())

    # ═══════════════════════════════════════
    # DETAILED ITINERARY
    # ═══════════════════════════════════════
    story.append(Paragraph(tr.t('export.detailedItinerary', 'Detailed Itinerary'), h1_style))
    story.append(HRFlowable(width='100%', thickness=1,
                              color=colors.HexColor('#2E4057')))

    for i, s in enumerate(stops):
        color = STOP_COLORS[i % len(STOP_COLORS)]
        heading_text = (f'<font color="{color}"><b>{i+1}.</b></font> '
                        f'<font name="Helvetica-Bold">{s.name}</font>')
        story.append(Paragraph(heading_text, stop_title_style))

        meta_parts = []
        if s.start_date:
            if s.end_date:
                meta_parts.append(f'{tr.fmt_date(s.start_date)} – {tr.fmt_date(s.end_date)}')
            else:
                meta_parts.append(tr.fmt_date(s.start_date))
        n = _nights(s)
        if n is not None:
            meta_parts.append(tr.nights_label(n))
        if s.distance_km:
            meta_parts.append(
                f'{s.distance_km:.0f} {tr.t("export.kmFromPrevious", "km from previous")}')
        if meta_parts:
            story.append(Paragraph(' · '.join(meta_parts), meta_style))

        if s.address:
            story.append(Paragraph(s.address, meta_style))
        if s.description:
            story.append(Paragraph(s.description, body_style))
        if s.url:
            story.append(Paragraph(s.url, url_style))

        if s.activities:
            story.append(Paragraph(
                f'<b>{tr.t("export.activities", "Activities")}:</b>', body_style))
            for act in s.activities:
                text = f'• <b>{act.name}</b>'
                if act.description:
                    text += f' — {act.description}'
                story.append(Paragraph(text, bullet_style))
                if act.url:
                    story.append(Paragraph(
                        f'&nbsp;&nbsp;<font color="#4285F4">{act.url}</font>',
                        ParagraphStyle('AU', fontSize=8, leftIndent=14, spaceAfter=2)))

        story.append(Spacer(1, 4))

    # ---- Bookmarks ----
    if bookmarks:
        story.append(PageBreak())
        story.append(Paragraph(
            tr.t('export.bookmarksResources', 'Bookmarks &amp; Resources'), h1_style))
        story.append(HRFlowable(width='100%', thickness=1,
                                  color=colors.HexColor('#2E4057')))
        for bm in bookmarks:
            text = f'<b>{bm.description}:</b> ' if bm.description else ''
            text += f'<font color="#4285F4">{bm.url}</font>'
            story.append(Paragraph(text, bullet_style))

    # ═══════════════════════════════════════
    # MAP PAGE
    # ═══════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph(tr.t('export.routeMap', 'Route Map'), h1_style))
    story.append(HRFlowable(width='100%', thickness=1,
                              color=colors.HexColor('#2E4057')))
    story.append(Spacer(1, 0.3*cm))

    map_bytes = _get_static_map_image(stops, trip, api_key, width=800, height=520)
    if map_bytes:
        story.append(RLImage(BytesIO(map_bytes), width=W, height=W * 520/800))
    else:
        story.append(Paragraph(
            tr.t('export.mapUnavailable', 'Map unavailable — no coordinates found.'),
            meta_style))

    doc.build(story)
    buf.seek(0)
    return buf
